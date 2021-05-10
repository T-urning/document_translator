import docx
import json
import time
import yaml
import random
import requests

from cProfile import Profile    # 评估性能
from pstats import Stats
from hashlib import md5
from tqdm import tqdm
from urllib3.util.retry import Retry
from requests.adapters import HTTPAdapter

# 加载配置文件
config = yaml.load(open('config.yaml', mode='r', encoding='utf-8'))
# 全局变量
TRANSLATE_API = 'https://fanyi-api.baidu.com/api/trans/vip/translate'
APP_ID = str(config['API']['appid']) 
APP_KEY = config['API']['appkey']
SLEEP_TIME = 1

# Generate salt and sign
def make_md5(s, encoding='utf-8'):
    return md5(s.encode(encoding)).hexdigest()


def translate(text, source_lang=None, target_lang=None):
    ''' 
    调用百度翻译 API 将文本 text 从 source_lang 语言翻译到 target_lang 语言，API 网址为：https://fanyi-api.baidu.com/product/113
    '''
    text = text.strip()
    if not text:
        return text
    if source_lang is None:
        source_lang = 'zh' # 'auto' -> 自动检测语种
    if target_lang is None:
        target_lang = 'en'
    salt = random.randint(32768, 65536)
    sign = make_md5(APP_ID + text + str(salt) + APP_KEY)

    # Build request
    headers = {'Content-Type': 'application/x-www-form-urlencoded'}
    payload = {'appid': APP_ID, 'q': text, 'from': source_lang, 'to': target_lang, 'salt': salt, 'sign': sign}
    
    s = requests.Session()
    retries = Retry(total=10)
    s.mount('https://', HTTPAdapter(max_retries=retries))
    # Send request
    r = s.post(TRANSLATE_API, params=payload, headers=headers)
    result = r.json()
    # 拦截错误
    if 'error_code' in result.keys():
        error_code = result['error_code']
        raise Exception(f'调用百度翻译 API 出现错误，错误码为 {error_code}，' +
            '请根据错误码到 https://fanyi-api.baidu.com/product/113 找相应的解决方法。')
    trans_list = [trans_res['dst'] for trans_res in result['trans_result']]
    return trans_list

def paras_generator(all_paras, batch_size=5):
    all_paras_new = [para for para in all_paras if len(para.text.strip()) > 0]
    total = len(all_paras_new)
    start = 0
    for start in range(0, total, batch_size):
        end = start + batch_size
        end = total if end > total else end
        yield all_paras_new[start: end]

def main():
    # 待翻译文档路径
    file_path = config['FILE']['file_path'] #'test_document.docx'
    
    # 读取文档
    doc = docx.Document(file_path)
    all_paras = doc.paragraphs  # 文档所有段落（标题+正文+图表题注）
    all_tables = doc.tables # 所有表格

    # 翻译所有段落，包括图片与表格的标题名
    print('处理所有段落，包括图片与表格的标题名...')
    batch_size = 10 # 为了提升翻译速度，每次传入 batch_size 个段落同时翻译
    for para_batch in tqdm(paras_generator(all_paras, batch_size)):
        # paras_generator 确保了选取每个 para 的文本都不为空
        para_combined = '\n'.join([para.text.strip() for para in para_batch])
        try:
            para_translated = translate(para_combined)
        except Exception as e:
            print('翻译失败')
            raise e
        else:   # 翻译成功则执行：
            assert len(para_batch) == len(para_translated)
            for para, text_target in zip(para_batch, para_translated):
                text_source = para.text.strip()
                if len(para.runs) < 1:
                    continue
                # 记录该段落的字体样式
                style = para.runs[0].style 
                para.text = para.text.replace(text_source, text_target)
                # 设置回原先的样式
                for run in para.runs:
                    run.style = style
            # 受翻译API访问频率的限制
            time.sleep(SLEEP_TIME)
    
    # 翻译表格内容
    print('处理所有表格...')
    for table in tqdm(all_tables): # 所有表格
        table_values = []
        cells_to_translate = []
        for row in table.rows: # 表格的所有行
            for cell in row.cells: # 该行的所有元素
                text_source = cell.text.strip()
                if len(text_source) > 0:
                    table_values.append(text_source)
                    cells_to_translate.append(cell)
        if table_values:
            table_values = '\n'.join(table_values)
            try:
                table_translated = translate(table_values)
            except Exception as e:
                print('翻译出错')
                raise e
            else: # 翻译成功
                assert len(cells_to_translate) == len(table_translated)
                for cell, text_target in zip(cells_to_translate, table_translated):
                    text_source = cell.text.strip()
                    cell.text = cell.text.replace(text_source, text_target)
                time.sleep(SLEEP_TIME)
    # 保存翻译结果              
    save_path = config['FILE']['save_path']             
    doc.save(save_path)
    print('翻译结果已保存。')

if __name__ == '__main__':
    # profiler = Profile()
    # profiler.runcall(main)
    # stats = Stats(profiler)
    # stats.sort_stats('cumulative')
    # stats.print_stats(50)
    main()

    

